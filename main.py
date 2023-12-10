from docx.api import Document

datasetPath = "path/to/doc/word/file.docx"

document = Document(datasetPath)

allTables = document.tables

data_to_print = ""
col_separator = ";"
output_index = 0

for single_table in document.tables:
    # Browse table rows
    table_rows = single_table.rows
    # Browse cells in the row
    row_cells = None
    for table_row in table_rows:
      row_cells = table_row.cells # list of cells
      for cell in row_cells:
        data_to_print += cell.text + col_separator
      # print the end_of_line
      data_to_print += '\n'
      print(data_to_print)

 # Sauvegarde du texte dans un fichier texte
 with open(f'out/dataTab{output_index}.csv', 'w', encoding='utf-8') as txt_file:
       txt_file.write(data_to_print)
 # reset the data file
 data_to_print = ""
 # go to next table_index
 output_index += 1



# Now zip the output to download the csv files
import os
import zipfile

def zipdir(path, ziph):
    # ziph is zipfile handle
    for root, dirs, files in os.walk(path):
        for file in files:
            ziph.write(os.path.join(root, file),
                       os.path.relpath(os.path.join(root, file),
                                       os.path.join(path, '..')))


with zipfile.ZipFile('out.zip', 'w', zipfile.ZIP_DEFLATED) as zipf:
    zipdir('out/', zipf)

